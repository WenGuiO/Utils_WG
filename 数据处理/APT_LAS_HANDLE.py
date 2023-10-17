# coding=utf-8
"""
Improvement, Commented by WenGui
"""
import os
import sys
import re
from collections import Counter
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

map_level = {'7': '高风险', '8': '高风险', '9': '高风险', '10': '高风险',
             '4': '中风险', '5': '中风险', '6': '中风险',
             '1': '低风险', '2': '低风险', '3': '低风险'}

event_len = 0
device = "APT"

# 可修改的参数
zone = "国际区"

def get_data(titles, eventName):
    """
    提前获取部分 统计 信息
    :param titles: A-Z字符串,用于遍历数据列
    :param eventName: 统计的指定 列 数据
    :return: 指定列所有数据, 所有事件类型, 每个事件的数量统计
    """
    total_event = []  # 所有事件
    types = []  # 唯一事件, 记录所有事件类型
    # 遍历每一列, 判断当前数据是属于 LAS 还是 APT
    # 判断完成类型处理数据即退出
    for title in titles:
        tmp = sheet1[title]
        # 检索 指定列 的数据(根据第一个 数据名称 来确认)
        if tmp[0].value == eventName:
            # 从第 2 行有效数据开始
            for row in tmp[1:]:  # ip
                # 对 LAS 数据的特殊操作
                if eventName == '事件数量(eventCount)':
                    ppp = int(row.value.replace('"', '').strip())
                    total_event.append(ppp)
                else:  # 其他
                    total_event.append(row.value.split("<br/>")[0])
                    # 确认是否在 types 添加新类型
                    if row.value in types:
                        continue
                    types.append(row.value.split("<br/>")[0])
            break

    # 统计各个事件类型的出现次数
    statistics = Counter(total_event)
    # 返回 指定列所有数据, 所有事件类型, 每个事件的数量统计
    return total_event, types, statistics


def handle_date_apt(wb, titles, events, event_type, event_counter, eventNames):
    """
    APT 数据处理
    :param wb: 工作表
    :param titles: A-Z字符串,用于遍历数据列
    :param events: 所有事件的 “事件类型” 列表
    :param event_type: 所有事件类型
    :param event_counter: 各事件类型的数量统计
    :param eventNames: 检索名
    :return: null
    """
    all_event = [1 for i in range(len(eventNames) + 1)]
    all_event[0] = events

    i = 0
    # 遍历处理 ['客户端IP', '服务端IP', '风险等级', '域名'] 列的数据
    for eventName in eventNames:
        i = i + 1
        # md1 存储的是: 如 “客户端IP” 这一列的所有数据
        md1 = []
        # 遍历 A-Z 列
        for c in titles:
            tmp = sheet1[c]  # 逐列
            # 寻找 ['客户端IP', '服务端IP', '风险等级', '域名'] 列
            if tmp[0].value == eventName:
                # 遍历列数据, 第 2 行数据开始
                for row in tmp[1:]:
                    md1.append(row.value)
        all_event[i] = md1
    # maps 存储所有的事件的关联数据: ['事件名称', '客户端IP', '服务端IP', '平台判断', '人工研判', '备注']
    maps = []  # 建立关联
    levels = []  # 存储风险等级，用于后续统计
    lens = len(events)  # 数据个数
    # 遍历所有事件, 处理 “域名” 列的数据
    # 并建立关联
    for i in range(lens):
        domain0 = all_event[4][i]  # 每一个事件的 “域名” 项
        if not domain0:
            domain0 = None
        else:
            # 域名为 ip
            if re.findall(r'\d+\.\d+\.\d+\.\d+', domain0) is not None:
                domain0 = None

        # 建立关联： 事件名称 + 源IP + 目的IP + 平台判定 + 人工判断 + 域名备注
        j = (all_event[0][i], all_event[1][i], all_event[2][i], all_event[3][i], '低风险', domain0)
        # 添加进 风险统计清单
        levels.append(all_event[3][i])
        maps.append(j)

    # 统计, 对于重复的数据会合并处理
    statistics = Counter(maps)
    # classes 存储所需创建的工作表 (以事件类型命名)
    classes = list(set([ev.split('【')[0] for ev in event_type]))
    # total_event 存储各个事件类型 数量统计信息
    total_event = [(i, event_counter[i]) for i in event_counter.keys()]

    wb1 = None
    # 创建 各个事件分类统计 工作表
    for cls in classes:
        try:
            wb1 = wb.create_sheet(title=cls.strip())
            print(f"正在处理表 {cls.strip()}", end=" ")
        except Exception as e:
            print(f"创建表!失败! \n<原因> " + str(e))
            continue
        # 添加首行说明
        wb1.append(['事件名称', '数量', '客户端IP', '服务端IP', '平台判断', '人工研判', '备注'])
        # 遍历 所有类型数据
        for et in event_type:
            # 如果属于当前 事件类型工作表
            if et.startswith(cls):
                for i in statistics.keys():
                    # 事件类型 匹配
                    if i[0] == et:
                        # 对 事件名称 的特殊处理, 因为需要匹配故只能在最后处理
                        en = i[0].split('】')
                        en = en[1] if en[1] != '' else en[0].split('【')[0]
                        # 逐行填入数据 ['事件名称', '数量', '客户端IP', '服务端IP', '平台判断', '人工研判', '备注']
                        wb1.append([en, statistics[i], i[1], i[2], i[3], i[4], i[5]])
        print("[已完成]")
    # 创建 数据统计 工作表
    print("\n# [INFO] 统计结果如下:")
    wb1 = wb.create_sheet(title='数据统计')
    wb1.append(['数据总量', str(lens)])
    print('# 数据总量: ', str(lens))
    wb1.append(['事件统计:'])
    print('\n# 事件统计:')
    for a, b in total_event:
        wb1.append([a, b])
        print(a, b)
    wb1.append(['风险等级统计:'])
    print('\n# 风险等级统计:')
    c_level = Counter(levels)
    total_lv = [(i, c_level[i]) for i in c_level.keys()]
    for a, b in total_lv:
        wb1.append([a, b])
        print(a, b)
    NEW_FILE_NAME = excel_file
    wb.save(NEW_FILE_NAME)


def handle_date_las(wb, titles, events, cnt, event_type, eventNames):
    client_ips = set()
    server_ips = set()
    all_event = [1 for i in range(len(eventNames) + 1)]  # 所有事件
    all_event[0] = events

    i = 0
    for eventName in eventNames:
        # md1 存储的是: 如 “客户端IP” 这一列的所有数据
        md1 = []
        # 遍历 A-Z 列
        for title in titles:
            tmp = sheet1[title]  # 逐列
            # 寻找 ['来源地址(srcAddress)', '目的地址(destAddress)', '事件级别(severity)', '原始事件(rawEvent)'] 列
            if tmp[0].value == eventName:
                # 遍历列数据, 第 2 行数据开始
                for row in tmp[1:]:
                    md1.append(row.value)
        i += 1
        all_event[i] = md1

    # maps 存储所有的事件的关联数据: ['事件名称', '客户端IP', '服务端IP', '平台判断', '人工研判', '备注']
    maps = []  # 建立关联
    levels = []  # 等级
    lens = len(events)  # 存储风险等级，用于后续统计
    event_ss = []   # 存储各个事件的 "事件名称" 项，用于各类事件的 数量统计
    global event_len   # 事件总数

    # 遍历所有事件, 并建立关联
    for i in range(lens):
        if '日志重复生成' in events[i]:
            continue
        if '指定接口光模块异常' in events[i]:
            continue
        event_len += cnt[i]
        # “域名” 列数据处理
        domain0 = all_event[4][i]
        if not domain0:
            domain0 = None
        else:
            ttt = re.findall('域名: (.+?)~~~', all_event[4][i])
            domain0 = ttt[0] if ttt else None

        # 建立关联： 事件名称 + 源IP + 目的IP + 平台判定 + 人工判断 + 域名备注
        j = (all_event[0][i], all_event[1][i], all_event[2][i], map_level.get(all_event[3][i].strip()), '低风险', domain0)

        # 添加 该事件数量 的事件
        for k in range(cnt[i]):
            maps.append(j)
            event_ss.append(all_event[0][i])
            levels.append(map_level.get(all_event[3][i].strip()))

    # 统计, 对于重复的数据会合并处理
    statistics = Counter(maps)
    # classes 存储所需创建的工作表 (以事件类型命名)
    classes = list(set([ev.split('：')[0] for ev in event_type]))
    ev22 = Counter(event_ss)
    totalEVENT = [(i, ev22[i]) for i in ev22.keys()]

    # 创建 各个事件分类统计 工作表
    print("# [INFO] 正在分类创建工作表...")
    for cls in classes:
        if '日志重复生成' in cls:
            continue
        if '指定接口光模块异常' in cls:
            continue
        # 对 工作表 名称的处理
        title = cls.replace('：', '-')[:-1] if cls.endswith('：') else cls.replace('：', '-')
        title = title.replace("\\", "-").replace("/", "-").replace("*", "-").replace("?", "-").replace("[", "-")\
            .replace("]", "-").strip()
        if len(title) >= 30:
            title = title[0:30]
        wb1 = None
        try:
            wb1 = wb.create_sheet(title=title)
            print(f"正在处理表 {title}", end=" ")
        except Exception as e:
            print(f"创建表!失败! \n<原因> " + str(e))
            continue
        wb1.append(['事件名称', '数量', '客户端IP', '服务端IP', '平台判断', '人工研判', '备注'])
        for et in event_type:
            # 如果属于当前 事件类型工作表
            if et.startswith(cls):
                for i in statistics.keys():
                    try:
                        if i[0] == et:
                            event_name = i[0].strip()
                            if event_name.endswith('：'):
                                event_name = event_name[:-1]
                            # "客户端", "服务端" 的处理
                            client_ip = i[1].strip() if i[1] else None
                            server_ip = i[2].strip() if i[2] else None
                            if client_ip:
                                client_ips.add(client_ip)
                            if server_ip:
                                server_ips.add(server_ip)
                            # 填充数据: ['事件名称', '数量', '客户端IP', '服务端IP', '平台判断', '人工研判', '备注']
                            wb1.append([event_name, statistics[i], client_ip, server_ip, i[3], i[4], i[5]])
                    except Exception as e:
                        print("\n数据处理错误:" + str(i))
                        print("<原因> " + str(e))
        print("[已完成]")

    # 创建 数据统计 工作表
    print("\n# [INFO] 统计结果如下:")
    wb1 = wb.create_sheet(title='数据统计')
    wb1.append(['数据总量', str(event_len)])
    print('# 数据总量: ', str(event_len))
    wb1.append(['事件统计:'])
    print('\n# 事件统计:')
    for a0, b in totalEVENT:
        a = a0.strip()
        if a.endswith('：'):
            a = a[:-1]
        wb1.append([a, b])
        print(a, b)

    wb1.append(['风险等级统计:'])
    print('\n# 风险等级统计:')
    c_level = Counter(levels)
    total_lv = [(i, c_level[i]) for i in c_level.keys()]
    for a, b in total_lv:
        wb1.append([a, b])
        print(a, b)
    wb1.append(['IP统计:'])
    print('\n# IP统计:')
    wb1.append(['客户端IP:', str(len(client_ips))])
    print('客户端IP:', str(len(client_ips)))
    wb1.append(['服务端IP:', str(len(server_ips))])
    print('服务端IP:', str(len(server_ips)))

    NEW_FILE_NAME = excel_file
    wb.save(NEW_FILE_NAME)


def read_data(originFile):
    """
    读取文件并处理
    :param originFile: EXCEL 文件
    :return:
    """
    global event_len
    global device
    # 打开 Excel 文件
    wb = load_workbook(originFile)
    # 获取所有工作表
    sheets = wb.worksheets
    # 获取第一张sheet
    global sheet1
    sheet1 = sheets[0]

    # 列名, 用于遍历数据列
    titles = 'ABCDEFGHIJKLMNOPQRSTUVWXYZ'

    # 判断 日志类型: APT or LAS
    if 'las' not in originFile and 'event' not in originFile:  # APT file
        # 获取: 所有事件、 所有事件类型、 每个事件数量统计
        events, event_type, event_counter = get_data(titles, '事件名称')
        event_len = len(events)
        handle_date_apt(wb, titles, events, event_type, event_counter, ['客户端IP', '服务端IP', '风险等级', '域名'])

    else:  # LAS file
        # 获取: 所有事件、 所有事件类型、 每个事件数量统计
        events, event_type, temp1 = get_data(titles, '事件名称(name)')
        # cnt 获取 事件数量
        event_counter, temp1, temp2 = get_data(titles, '事件数量(eventCount)')
        device = "LAS"
        handle_date_las(wb, titles, events, event_counter, event_type,
                        ['来源地址(srcAddress)', '目的地址(destAddress)', '事件级别(severity)', '原始事件(rawEvent)'])


def generate_docx(origin_file):
    # 打开 Excel 文件
    workbook = load_workbook(origin_file)
    docx_file = f'{origin_file}.docx'

    # 创建 Word 文档
    doc = Document()

    sheet_all = []  # 事件类型统计
    sheet_total = []    # 事件数量统计
    total = 0
    # 遍历每个表格，但跳过最后一个表格
    for index, sheet_name in enumerate(workbook.sheetnames):
        try:
            # 第一个表格和最后一个表，跳过
            if index == 0 or index == len(workbook.sheetnames) - 1:
                continue
            # 进度
            print(f'# 正在处理第 {index} 个表格  ', end='')

            worksheet = workbook[sheet_name]

            # 统计 事件数量
            for row in worksheet.iter_rows(min_row=2, min_col=2):
                if row[0].value is None:
                    break
                total += int(row[0].value)

            # 添加表格标题
            # 添加表格标题，黑色字体，字体大小为16
            title_paragraph = doc.add_paragraph(f'{index}、{sheet_name}', style='Heading 3')
            title_paragraph.style.font.name = '宋体'
            title_paragraph.style.font.color.rgb = RGBColor(0, 0, 0)  # 设置字体颜色
            title_paragraph.style.font.size = Pt(16)

            # 小标题1
            title_paragraph = doc.add_paragraph()
            title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run2 = title_paragraph.add_run('（1）风险事件记录')
            run2.bold = True
            run2.font.name = '仿宋_GB2312'
            run2._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            run2.font.size = Pt(16)

            # 添加表格
            table = doc.add_table(rows=worksheet.max_row, cols=worksheet.max_column, style='Table Grid')
            for row in table.rows:
                for cell in row.cells:
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    cell.paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER

            # 设置表格的宽度和每一列的宽度
            table.width = Inches(8)
            column_width = int(table.width / 7)
            for i in range(7):
                table.columns[i].width = column_width

            # 遍历每一行
            for row_index, row in enumerate(worksheet.iter_rows()):
                # 遍历每一个单元格
                for column_index, cell in enumerate(row):
                    # 复制单元格内容到 Word 表格中
                    cell_copy = table.cell(row_index, column_index).paragraphs[0].add_run(str(cell.value))
                    # 设置单元格内容的字体
                    cell_copy.font.name = '宋体'
                    cell_copy.font.size = Pt(11)
                    cell_copy.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    # 设置单元格边框颜色
                    cell_border = cell.border
                    if cell_border.top.style is not None and cell_border.top.color is not None:
                        table.cell(row_index, column_index).top.border.color.rgb = cell_border.top.color.rgb
                        table.cell(row_index, column_index).top.border.size = Pt(1.5)
                    if cell_border.bottom.style is not None and cell_border.bottom.color is not None:
                        table.cell(row_index, column_index).bottom.border.color.rgb = cell_border.bottom.color.rgb
                        table.cell(row_index, column_index).bottom.border.size = Pt(1.5)
                    if cell_border.left.style is not None and cell_border.left.color is not None:
                        table.cell(row_index, column_index).left.border.color.rgb = cell_border.left.color.rgb
                        table.cell(row_index, column_index).left.border.size = Pt(1.5)
                    if cell_border.right.style is not None and cell_border.right.color is not None:
                        table.cell(row_index, column_index).right.border.color.rgb = cell_border.right.color.rgb
                        table.cell(row_index, column_index).right.border.size = Pt(1.5)

            sheet_all.append(sheet_name)
            sheet_total.append(total)
        except:
            print(F'第{index}个表格 {sheet_name}  [失败]')
        # 添加空白段落
        # doc.add_paragraph()
        # 小标题2
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run1 = title_paragraph.add_run('（2）风险事件小结')
        run1.bold = True
        run1.font.name = '仿宋_GB2312'
        run1._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        run1.font.size = Pt(16)
        # 添加结论
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = title_paragraph.add_run(f'    本时段内，杭州亚运村{zone}网络流量中发现{sheet_name}类攻击{total}次。')
        run.bold = False
        run.font.name = '仿宋_GB2312'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        run.font.size = Pt(16)

        print(f'[已完成]')
        # 添加空白段落作为表格与表格之间的间隔
        # doc.add_paragraph()

    text = f"    通过对{device}平台告警数据分析，本次发现告警" + str(event_len) + "条" + "，其中"
    for sheetName, total in zip(sheet_all, sheet_total):
        text = text + f'{sheetName}类{total}条，'
    print('\n' + text + "详细信息如下。")

    # 以下为添加 汇总至文档末尾
    # 标题
    title_paragraph = doc.add_paragraph(f'总结', style='Heading 3')
    title_paragraph.style.font.name = '宋体'
    title_paragraph.style.font.color.rgb = RGBColor(0, 0, 0)
    title_paragraph.style.font.size = Pt(16)

    run = doc.add_paragraph().add_run(f'{text}详细信息如下。')
    run.bold = False
    run.font.name = '仿宋'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    run.font.size = Pt(16)

    # 关闭execl
    workbook.close()
    # 保存 Word 文件
    doc.save(docx_file)

    # 再重新整理格式
    doc = Document(docx_file)
    # 设置表格第一行底纹为灰色
    for table in doc.tables:
        for idx, cell in enumerate(table.rows[0].cells):
            shading_element = OxmlElement('w:shd')
            shading_element.set(qn('w:fill'), 'C0C0C0')
            cell._element.get_or_add_tcPr().append(shading_element)
    # 遍历文档中的所有表格
    for table in doc.tables:
        # 遍历表格中的所有单元格
        for row in table.rows:
            for cell in row.cells:
                # 如果单元格的内容为'None'，则将其替换为空
                if cell.text == 'None':
                    cell.text = ''
    doc.save(docx_file)


if __name__ == '__main__':
    if len(sys.argv) != 2:
        print("Usage: python APT_LAS_HANDLE.py 文件路径")
        exit(1)

    excel_file = sys.argv[1]
    if os.path.exists(excel_file):
        # 处理 excel
        print("# [INFO] 正在处理数据...")
        read_data(excel_file)
        print('\n# [INFO] 已完成 原EXCEL文件 的数据处理')

        # 生成 docx
        print("\n# [INFO] 正在生成数据Word文档...")
        generate_docx(excel_file)
    else:
        print("错误: 文件不存在")
        print("Usage: python APT_LAS_HANDLE.py 文件路径")
